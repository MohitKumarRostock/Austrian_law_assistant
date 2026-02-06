"""
kahm_dashboard_gradio.py

Professional dashboard for KAHM embedding evaluation and interactive retrieval.

What it does
------------
1) Reads the publication-ready markdown report produced by evaluate_three_embeddings.py
   and visualizes key results via tables and charts.
2) Loads the 200 human-labeled test queries and supports:
     - selecting a test query and retrieving Top-K sentences using:
         (a) KAHM(query→MB corpus)  : KAHM-regressed query embedding searched on MB corpus index
         (b) Full-KAHM              : KAHM-regressed query embedding searched on KAHM-transformed corpus index
     - entering an arbitrary Austrian-law-related query and retrieving Top-K results.
   The dashboard emphasizes *sentence-level evidence* (possibly spanning multiple law topics),
   not single-label classification.
3) Provides a short "Science behind KAHM" explanation with citations (see KAHM science tab).
4) Presents a management-ready funding pitch (see Funding pitch tab).

How to run
----------
pip install gradio faiss-cpu pandas numpy plotly joblib pyarrow
python kahm_dashboard_gradio.py

If you run from your project root, default relative paths should work.

Notes
-----
- The dashboard intentionally reuses the same loading + KAHM regression helpers from
  evaluate_three_embeddings.py to ensure consistency.
- If you changed filenames/paths, set them in the sidebar (left).
"""
from __future__ import annotations

import re
import sys
import os
from dataclasses import dataclass
from functools import lru_cache
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

import numpy as np
import pandas as pd

import gradio as gr
import plotly.graph_objects as go

# Optional dependencies for LLM post-processing (summarization / plain-language interpretation)
# The dashboard remains fully functional without these dependencies.
try:  # Transformers seq2seq summarizers (CPU-friendly)
    import torch  # type: ignore
    from transformers import AutoModelForSeq2SeqLM, AutoTokenizer  # type: ignore
except Exception:  # pragma: no cover
    torch = None  # type: ignore
    AutoModelForSeq2SeqLM = None  # type: ignore
    AutoTokenizer = None  # type: ignore

try:  # llama.cpp bindings (local GGUF inference)
    from llama_cpp import Llama  # type: ignore
except Exception:  # pragma: no cover
    Llama = None  # type: ignore

# Optional dependency: faiss
try:
    import faiss  # type: ignore
except Exception as e:  # pragma: no cover
    faiss = None  # type: ignore


# ----------------------------- Project imports -----------------------------
# Ensure we can import evaluate_three_embeddings.py when running from other working dirs.
_THIS_DIR = Path(__file__).resolve().parent
if str(_THIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THIS_DIR))

try:
    import evaluate_three_embeddings as eval3  # noqa: E402
except Exception as e:  # pragma: no cover
    eval3 = None  # type: ignore
    _EVAL3_IMPORT_ERROR = e
else:
    _EVAL3_IMPORT_ERROR = None


# ----------------------------- Parsing helpers -----------------------------
_NUM_CI_RE = re.compile(
    r"^\s*([+-]?\d+(?:\.\d+)?)\s*\(\s*([+-]?\d+(?:\.\d+)?)\s*,\s*([+-]?\d+(?:\.\d+)?)\s*\)\s*$"
)


@dataclass(frozen=True)
class MeanCI:
    mean: float
    lo: float
    hi: float

    @staticmethod
    def parse(cell: str) -> Optional["MeanCI"]:
        m = _NUM_CI_RE.match(str(cell))
        if not m:
            return None
        return MeanCI(float(m.group(1)), float(m.group(2)), float(m.group(3)))


def _extract_all_md_tables(md: str, header_prefix: str) -> List[List[str]]:
    """
    Extract all markdown tables whose header row begins with `header_prefix`.
    Each table is returned as a list of lines (including header + separator + rows).
    """
    lines = md.splitlines()
    tables: List[List[str]] = []
    i = 0
    while i < len(lines):
        if lines[i].strip().startswith(header_prefix):
            out: List[str] = []
            j = i
            while j < len(lines) and lines[j].strip().startswith('|'):
                out.append(lines[j].rstrip())
                j += 1
            if len(out) >= 3:
                tables.append(out)
            i = j
            continue
        i += 1
    return tables


def _md_table_to_df(table_lines: Sequence[str]) -> pd.DataFrame:
    """
    Convert markdown table lines into a DataFrame (string cells preserved).
    """
    # Header + rows; skip the separator row (| --- | --- |)
    header = [c.strip() for c in table_lines[0].strip().strip("|").split("|")]
    rows: List[List[str]] = []
    for line in table_lines[2:]:
        parts = [c.strip() for c in line.strip().strip("|").split("|")]
        if len(parts) != len(header):
            # tolerate ragged rows by padding/truncation
            if len(parts) < len(header):
                parts = parts + [""] * (len(header) - len(parts))
            else:
                parts = parts[: len(header)]
        rows.append(parts)
    return pd.DataFrame(rows, columns=header)


def parse_report_tables(report_path: str) -> Dict[str, pd.DataFrame]:
    """
    Parse key tables from the markdown report produced by evaluate_three_embeddings.py.

    Notes
    -----
    The report contains multiple markdown tables with similar headers (e.g., paired-delta tables).
    We therefore extract *all* matching tables and then map them to expected ids by order.

    Returns
    -------
    Dict[str, DataFrame]
        Keys: table1, table2, table3, table5, table9 (when found).
    """
    p = Path(report_path)
    if not p.exists():
        raise FileNotFoundError(f"Report not found: {report_path}")
    md = p.read_text(encoding="utf-8", errors="replace")

    tables: Dict[str, pd.DataFrame] = {}

    # Table 1 (main metrics)
    t1_all = _extract_all_md_tables(md, "| Method | Hit@")
    if t1_all:
        tables["table1"] = _md_table_to_df(t1_all[0])

    # Paired delta tables (Table 2 and Table 3 share the same header prefix)
    delta_all = _extract_all_md_tables(md, "| Metric | Δ (KAHM")
    if len(delta_all) >= 1:
        tables["table2"] = _md_table_to_df(delta_all[0])
    if len(delta_all) >= 2:
        tables["table3"] = _md_table_to_df(delta_all[1])

    # Table 5 (routing sweep)
    t5_all = _extract_all_md_tables(md, "| τ | Coverage (KAHM)")
    if t5_all:
        tables["table5"] = _md_table_to_df(t5_all[0])

    # Table 9 (alignment)
    t9_all = _extract_all_md_tables(md, "| Quantity | Estimate")
    if t9_all:
        tables["table9"] = _md_table_to_df(t9_all[0])

    return tables


def _metric_table_long(df_table1: pd.DataFrame, metrics: List[str]) -> pd.DataFrame:
    """
    Convert Table 1 into long format with mean/ci for plotting.
    metrics: column names in df_table1
    """
    out_rows: List[Dict[str, Any]] = []
    for _, row in df_table1.iterrows():
        method = str(row.get("Method", "")).strip()
        for m in metrics:
            cell = str(row.get(m, "")).strip()
            parsed = MeanCI.parse(cell)
            if parsed is None:
                # fallback: try to parse the first float
                try:
                    mean = float(re.findall(r"[+-]?\d+(?:\.\d+)?", cell)[0])
                except Exception:
                    continue
                parsed = MeanCI(mean=mean, lo=np.nan, hi=np.nan)
            out_rows.append(
                {"Method": method, "Metric": m, "Mean": parsed.mean, "CI_lo": parsed.lo, "CI_hi": parsed.hi}
            )
    return pd.DataFrame(out_rows)


def fig_main_metrics(df_table1: pd.DataFrame) -> go.Figure:
    metrics = [
        "Hit@10",
        "MRR@10 (unique laws)",
        "Top-1 accuracy",
        "Majority-vote accuracy (predominance ≥ 0.50)",
    ]
    long = _metric_table_long(df_table1, metrics)
    fig = go.Figure()

    for method in long["Method"].unique():
        sub = long[long["Method"] == method]
        y = sub["Mean"].to_numpy(dtype=float)
        hi = sub["CI_hi"].to_numpy(dtype=float)
        lo = sub["CI_lo"].to_numpy(dtype=float)
        arr = np.maximum(0.0, hi - y)
        arrm = np.maximum(0.0, y - lo)
        finite = np.isfinite(arr).all() and np.isfinite(arrm).all()

        fig.add_trace(
            go.Bar(
                name=method,
                x=sub["Metric"],
                y=y,
                error_y=dict(
                    type="data",
                    array=arr,
                    arrayminus=arrm,
                    visible=bool(finite),
                ),
            )
        )

    fig.update_layout(
        barmode="group",
        height=420,
        margin=dict(l=30, r=30, t=50, b=30),
        title="Retrieval quality metrics (mean ± 95% bootstrap CI)",
        yaxis_title="Score",
        legend_title="Method",
    )
    return fig


def _bar_metric(df_table1: pd.DataFrame, metric: str, *, title: str, yaxis_title: str) -> go.Figure:
    long = _metric_table_long(df_table1, [metric])
    fig = go.Figure()

    for method in long["Method"].unique():
        sub = long[long["Method"] == method]
        y = sub["Mean"].to_numpy(dtype=float)
        hi = sub["CI_hi"].to_numpy(dtype=float)
        lo = sub["CI_lo"].to_numpy(dtype=float)
        arr = np.maximum(0.0, hi - y)
        arrm = np.maximum(0.0, y - lo)
        finite = np.isfinite(arr).all() and np.isfinite(arrm).all()

        fig.add_trace(
            go.Bar(
                name=method,
                x=[metric],
                y=y,
                error_y=dict(
                    type="data",
                    array=arr,
                    arrayminus=arrm,
                    visible=bool(finite),
                ),
            )
        )

    fig.update_layout(
        barmode="group",
        height=380,
        margin=dict(l=30, r=30, t=50, b=30),
        title=title,
        yaxis_title=yaxis_title,
        legend_title="Method",
    )
    return fig


def fig_mean_consensus_fraction(df_table1: pd.DataFrame) -> go.Figure:
    return _bar_metric(
        df_table1,
        "Mean consensus fraction",
        title="Mean consensus fraction (mean ± 95% bootstrap CI)",
        yaxis_title="Fraction",
    )


def fig_mean_lift(df_table1: pd.DataFrame) -> go.Figure:
    return _bar_metric(
        df_table1,
        "Mean lift vs prior",
        title="Mean lift vs prior (mean ± 95% bootstrap CI)",
        yaxis_title="Lift",
    )


def fig_routing(df_table5: pd.DataFrame) -> go.Figure:
    required = [
        "τ",
        "Coverage (KAHM)",
        "Accuracy among covered (KAHM)",
        "Coverage (Mixedbread)",
        "Accuracy among covered (Mixedbread)",
    ]
    missing = [c for c in required if c not in df_table5.columns]
    if missing:
        fig = go.Figure()
        fig.update_layout(
            height=240,
            margin=dict(l=30, r=30, t=50, b=30),
            title="Vote-based routing tradeoff (table parsing incomplete)",
            annotations=[
                dict(
                    text=f"Missing columns in parsed Table 5: {', '.join(missing)}",
                    x=0.5,
                    y=0.5,
                    xref="paper",
                    yref="paper",
                    showarrow=False,
                )
            ],
        )
        return fig

    def _parse_col(col: str) -> Tuple[np.ndarray, np.ndarray, np.ndarray]:
        parsed = [MeanCI.parse(v) for v in df_table5[col].tolist()]
        means = np.array([p.mean if p else np.nan for p in parsed], dtype=float)
        lo = np.array([p.lo if p else np.nan for p in parsed], dtype=float)
        hi = np.array([p.hi if p else np.nan for p in parsed], dtype=float)
        return means, lo, hi

    tau = df_table5["τ"].astype(float).to_numpy()
    cov_k, _, _ = _parse_col("Coverage (KAHM)")
    acc_k, _, _ = _parse_col("Accuracy among covered (KAHM)")
    cov_m, _, _ = _parse_col("Coverage (Mixedbread)")
    acc_m, _, _ = _parse_col("Accuracy among covered (Mixedbread)")

    fig = go.Figure()
    fig.add_trace(go.Scatter(x=tau, y=acc_m, mode="lines+markers", name="Mixedbread: acc|covered"))
    fig.add_trace(go.Scatter(x=tau, y=acc_k, mode="lines+markers", name="KAHM: acc|covered"))
    fig.add_trace(go.Scatter(x=tau, y=cov_m, mode="lines+markers", name="Mixedbread: coverage", line=dict(dash="dash")))
    fig.add_trace(go.Scatter(x=tau, y=cov_k, mode="lines+markers", name="KAHM: coverage", line=dict(dash="dash")))

    fig.update_layout(
        height=420,
        margin=dict(l=30, r=30, t=50, b=30),
        title="Vote-based routing tradeoff (coverage vs accuracy among covered)",
        xaxis_title="Predominance threshold τ",
        yaxis_title="Probability",
        legend_title="Curve",
    )
    return fig


# ----------------------------- Retrieval engine -----------------------------
@dataclass
class CorpusIndex:
    name: str
    sentence_ids: np.ndarray  # shape (n,)
    embeddings: np.ndarray  # shape (n, d)
    faiss_index: Any  # faiss.Index


def _pick_first_existing(cols: Sequence[str], available: Sequence[str]) -> Optional[str]:
    aset = set(available)
    for c in cols:
        if c in aset:
            return c
    return None


@lru_cache(maxsize=4)
def _load_dataframe(corpus_parquet: str) -> pd.DataFrame:
    if eval3 is None:
        raise ImportError(f"Could not import evaluate_three_embeddings.py: {_EVAL3_IMPORT_ERROR}")
    df = eval3.load_corpus_parquet(corpus_parquet)
    # ensure sentence_id exists and is int64
    if "sentence_id" not in df.columns:
        raise KeyError("Corpus parquet must contain column 'sentence_id'")
    df["sentence_id"] = df["sentence_id"].astype(np.int64)
    if df["sentence_id"].duplicated().any():
        raise ValueError("Corpus parquet contains non-unique sentence_id values")
    return df


@lru_cache(maxsize=4)
def _load_bundle(npz_path: str) -> Dict[str, np.ndarray]:
    if eval3 is None:
        raise ImportError(f"Could not import evaluate_three_embeddings.py: {_EVAL3_IMPORT_ERROR}")
    return eval3.load_npz_bundle(npz_path)


def _subset_to_metadata(bundle: Dict[str, np.ndarray], df: pd.DataFrame) -> Tuple[np.ndarray, np.ndarray]:
    ids = bundle["sentence_ids"].astype(np.int64)
    emb = bundle["emb"].astype(np.float32)
    if emb.ndim != 2:
        raise ValueError("NPZ bundle embeddings must be a 2D array")
    if ids.ndim != 1:
        raise ValueError("NPZ bundle sentence_ids must be a 1D array")
    if emb.shape[0] != ids.shape[0]:
        raise ValueError(f"NPZ bundle mismatch: {emb.shape[0]=} vs {ids.shape[0]=}")

    meta_ids = df["sentence_id"].astype(np.int64).to_numpy()
    keep = np.isin(ids, meta_ids)
    return ids[keep], emb[keep]


def _build_faiss_index(embeddings: np.ndarray) -> Any:
    if faiss is None:
        raise ImportError("faiss is not installed. Install faiss-cpu (or faiss-gpu).")
    d = int(embeddings.shape[1])
    index: Any = faiss.IndexFlatIP(d)
    # Ensure contiguous float32
    X = np.ascontiguousarray(embeddings.astype(np.float32))
    index.add(X)
    return index


@lru_cache(maxsize=2)
def _load_idf_svd_model(path: str):
    if eval3 is None:
        raise ImportError(f"Could not import evaluate_three_embeddings.py: {_EVAL3_IMPORT_ERROR}")
    return eval3.load_idf_svd_model(path)


@lru_cache(maxsize=2)
def _load_kahm_model(path: str) -> dict:
    if eval3 is None:
        raise ImportError(f"Could not import evaluate_three_embeddings.py: {_EVAL3_IMPORT_ERROR}")
    return eval3.load_kahm_model(path)


def _embed_query_kahm(
    query: str,
    *,
    idf_svd_model_path: str,
    kahm_model_path: str,
    kahm_mode: str,
    kahm_batch: int,
) -> np.ndarray:
    if eval3 is None:
        raise ImportError(f"Could not import evaluate_three_embeddings.py: {_EVAL3_IMPORT_ERROR}")
    pipe = _load_idf_svd_model(idf_svd_model_path)
    X = eval3.embed_queries_idf_svd(pipe, [query])  # shape (1, d)
    model = _load_kahm_model(kahm_model_path)
    Y = eval3.kahm_regress_once(model, X, mode=kahm_mode, batch_size=kahm_batch, show_progress=False)  # (1, d)
    # Ensure cosine-similarity semantics for IndexFlatIP
    Y = eval3.l2_normalize_rows(Y.astype(np.float32))
    return Y[0]


def _search_topk(corpus: CorpusIndex, q_emb: np.ndarray, top_k: int) -> Tuple[np.ndarray, np.ndarray]:
    q = np.ascontiguousarray(q_emb.astype(np.float32).reshape(1, -1))
    D, I = corpus.faiss_index.search(q, int(top_k))
    scores = D[0]
    idxs = I[0]

    # FAISS may return -1 indices if k > n or on some index types.
    valid = (idxs >= 0) & (idxs < len(corpus.sentence_ids))
    if not np.all(valid):
        scores = scores[valid]
        idxs = idxs[valid]

    sids = corpus.sentence_ids[idxs]
    return scores, sids


def _format_hits(
    df_meta: pd.DataFrame,
    scores: np.ndarray,
    sentence_ids: np.ndarray,
) -> Tuple[pd.DataFrame, go.Figure]:
    # Prepare metadata mapping
    meta = df_meta.set_index("sentence_id", drop=False)

    available_cols = list(df_meta.columns)
    text_col = _pick_first_existing(
        ["sentence", "sentence_text", "text", "content", "passage", "segment"],
        available_cols,
    )
    page_col = _pick_first_existing(
        ["page_no", "page", "pageno", "page_number", "pdf_page"],
        available_cols,
    )
    law_col = _pick_first_existing(["law_type", "law", "law_id", "law_name"], available_cols) or "law_type"

    rows: List[Dict[str, Any]] = []
    law_types: List[str] = []
    for r, (sid, sc) in enumerate(zip(sentence_ids.tolist(), scores.tolist()), start=1):
        rec: Dict[str, Any] = {"rank": r, "score": float(round(sc, 4))}
        if int(sid) in meta.index:
            row = meta.loc[int(sid)]
            rec["law_type"] = str(row.get(law_col, ""))
            if page_col:
                rec["page_no"] = row.get(page_col, "")
            else:
                rec["page_no"] = ""
            if text_col:
                rec["sentence"] = str(row.get(text_col, ""))
            else:
                rec["sentence"] = ""
        else:
            rec["law_type"] = ""
            rec["page_no"] = ""
            rec["sentence"] = ""
        rows.append(rec)
        law_types.append(str(rec["law_type"]))

    out_df = pd.DataFrame(rows)
    # Law-type distribution
    ser = pd.Series(law_types).replace("", np.nan).dropna()
    counts = ser.value_counts().head(12)
    fig = go.Figure()
    fig.add_trace(go.Bar(x=counts.index.tolist(), y=counts.values.tolist(), name="count"))
    fig.update_layout(
        height=260,
        margin=dict(l=10, r=10, t=30, b=20),
        title="Top law-types in retrieved Top-K (informative only)",
        xaxis_title="law_type",
        yaxis_title="count",
    )
    return out_df, fig




# ----------------------------- LLM post-processing -----------------------------

_DEFAULT_SEQ2SEQ_MODEL = "deutsche-telekom/mt5-small-sum-de-mit-v1"


def _safe_str(x: object) -> str:
    try:
        return str(x)
    except Exception:
        return ""


def _hits_to_compact_text(
    df_hits: pd.DataFrame,
    *,
    max_items: int = 10,
    max_chars: int = 9000,
) -> str:
    """Create a compact, citation-friendly text block from a retrieval DataFrame."""
    if df_hits is None or df_hits.empty:
        return ""
    cols = set(df_hits.columns)
    need = {"rank", "law_type", "sentence"}
    if not need.issubset(cols):
        return ""

    rows = []
    for _, r in df_hits.head(int(max_items)).iterrows():
        rank = _safe_str(r.get("rank", "")).strip()
        law = _safe_str(r.get("law_type", "")).strip()
        sent = _safe_str(r.get("sentence", "")).strip()
        if not sent:
            continue
        # Keep the model input compact and consistent.
        if len(sent) > 800:
            sent = sent[:800].rstrip() + "…"
        prefix = f"[{rank}]" if rank else ""
        law_part = f" ({law})" if law else ""
        rows.append(f"{prefix}{law_part} {sent}")

    out = "\n".join(rows).strip()
    if len(out) > int(max_chars):
        out = out[: int(max_chars)].rstrip() + "…"
    return out



def _evidence_stats(df_hits: pd.DataFrame, *, max_items: int = 10) -> Dict[str, float]:
    """Compute simple evidence-quality stats to support grounded prompting."""
    if df_hits is None or df_hits.empty:
        return {"n_sent": 0.0, "n_law_types": 0.0, "law_type_ratio": 0.0}

    n_sent = 0
    law_types: List[str] = []
    for _, r in df_hits.head(int(max_items)).iterrows():
        sent = _safe_str(r.get("sentence", "")).strip()
        if not sent:
            continue
        n_sent += 1
        law = _safe_str(r.get("law_type", "")).strip()
        if law:
            law_types.append(law)

    n_law = float(len(set(law_types))) if law_types else 0.0
    ratio = float(n_law / n_sent) if n_sent else 0.0
    return {"n_sent": float(n_sent), "n_law_types": n_law, "law_type_ratio": ratio}


def _key_excerpts_md(df_hits: pd.DataFrame, *, max_items: int = 6) -> str:
    """Create a compact bullet list of key excerpts for transparency."""
    if df_hits is None or df_hits.empty:
        return ""

    lines: List[str] = []
    for _, r in df_hits.head(int(max_items)).iterrows():
        rank = _safe_str(r.get("rank", "")).strip()
        law = _safe_str(r.get("law_type", "")).strip()
        sent = _safe_str(r.get("sentence", "")).strip()
        if not sent:
            continue
        if len(sent) > 500:
            sent = sent[:500].rstrip() + "…"

        rank0 = f"[{rank}]" if rank else ""
        law0 = f"**{law}**: " if law else ""
        lines.append(f"- {rank0} {law0}{sent}")

    return "\n".join(lines).strip()


@lru_cache(maxsize=2)
def _load_seq2seq(model_id: str):
    if AutoTokenizer is None or AutoModelForSeq2SeqLM is None:
        raise ImportError("transformers/torch are not available. Install: pip install transformers torch")

    # Tokenizer policy (robust + warning-free for legal text):
    # - Force the *slow* tokenizer (SentencePiece/Python) to avoid fast-tokenizer conversion warnings,
    #   including byte-fallback limitations that can surface as <unk> for rare characters.
    # - Explicitly set T5's legacy behavior when supported to prevent the "legacy behaviour" warning.
    tok_kwargs = {"use_fast": False}
    try:
        tok = AutoTokenizer.from_pretrained(model_id, **tok_kwargs, legacy=True)
    except TypeError:
        # Non-T5 tokenizers (or older Transformers versions) may not accept `legacy`.
        tok = AutoTokenizer.from_pretrained(model_id, **tok_kwargs)
    except Exception:
        # Last-resort fallback: if a repository does not ship a slow tokenizer, load the default.
        # (This may reintroduce warnings, but ensures the dashboard remains functional.)
        tok = AutoTokenizer.from_pretrained(model_id)

    # Seq2seq-only backend (mT5/T5/BART/PEGASUS). If a causal model is provided, loading will fail.
    mdl = AutoModelForSeq2SeqLM.from_pretrained(model_id)
    mdl.eval()
    return tok, mdl



@lru_cache(maxsize=1)
def _load_llama_cpp(gguf_path: str, n_ctx: int, n_threads: int):
    if Llama is None:
        raise ImportError("llama_cpp is not available. Install: pip install llama-cpp-python")
    p = Path(gguf_path)
    if not p.exists():
        raise FileNotFoundError(f"GGUF model not found: {gguf_path}")
    n_threads0 = int(n_threads)
    if n_threads0 <= 0:
        # UI allows 0=auto; default to available logical CPUs.
        n_threads0 = max(1, (os.cpu_count() or 4))
    return Llama(model_path=str(p), n_ctx=int(n_ctx), n_threads=n_threads0, verbose=False)


def _seq2seq_summarize(
    text_in: str,
    *,
    model_id: str,
    max_new_tokens: int = 160,
) -> str:
    if torch is None:
        raise ImportError("transformers/torch are not available. Install: pip install transformers torch")
    tok, mdl = _load_seq2seq(model_id)
    # mT5-style models commonly use a prefix.
    prompt = f"summarize: {text_in}"
    inputs = tok(
        prompt,
        return_tensors="pt",
        truncation=True,
        max_length=896,
    )
    with torch.no_grad():
        out = mdl.generate(
            **inputs,
            max_new_tokens=int(max_new_tokens),
            num_beams=4,
            do_sample=False,
            no_repeat_ngram_size=3,
        )
    return tok.decode(out[0], skip_special_tokens=True).strip()


def _llama_cpp_generate(
    compact_hits: str,
    *,
    query_text: str,
    method_label: str,
    output_language: str,
    gguf_path: str,
    max_new_tokens: int,
    temperature: float,
    n_ctx: int,
    n_threads: int,
    task: str,
    strict_grounding: bool,
    evidence_stats: Dict[str, float],
    heterogeneity_threshold: int,
    min_evidence_sentences: int,
) -> str:
    """Generate grounded, citation-forward summaries/paraphrases via a local GGUF instruct model."""
    llm = _load_llama_cpp(gguf_path, n_ctx=n_ctx, n_threads=n_threads)

    # Language handling
    lang = (output_language or "German").strip().lower()
    if lang.startswith("en"):
        lang_hint = "Write in English."
        unsupported_msg = "Not supported by the provided citations."
        caveat = "Not legal advice."
    else:
        lang_hint = "Schreibe auf Deutsch."
        unsupported_msg = "Nicht aus den Zitaten ableitbar."
        caveat = "Keine Rechtsberatung."

    task0 = (task or "Summary + plain-language interpretation").strip().lower()
    wants_per_sentence = (
        "per-sentence" in task0
        or "per sentence" in task0
        or "per-satz" in task0
        or "paraphrase" in task0
        or "paraphr" in task0
    )
    per_sentence_only = wants_per_sentence and ("only" in task0 or "nur" in task0)

    n_sent = int(evidence_stats.get("n_sent", 0.0) or 0)
    n_law = int(evidence_stats.get("n_law_types", 0.0) or 0)
    heterogeneous = bool(heterogeneity_threshold and n_law >= int(heterogeneity_threshold))
    sparse = bool(min_evidence_sentences and n_sent < int(min_evidence_sentences))

    evidence_snapshot = f"{n_sent} retrieved sentences; {n_law} distinct law types."
    flags = []
    if sparse:
        flags.append("sparse")
    if heterogeneous:
        flags.append("heterogeneous")
    if flags:
        evidence_snapshot += " Evidence is " + ", ".join(flags) + "."

    # Task selection
    if per_sentence_only:
        task_hint = (
            "Paraphrase each retrieved sentence individually in plain, professional language. "
            "Keep each paraphrase to one short sentence. Do not add new facts."
        )
        output_format = f"""Output format (Markdown):
- Per-sentence paraphrases (bullets; each bullet starts with [rank])
- Caveat: one sentence stating this is not legal advice ({caveat})
"""
    elif task0.startswith("summary") and ("interpret" not in task0 and "interpretation" not in task0):
        task_hint = "Provide a concise professional summary of the retrieved sentences."
        output_format = f"""Output format (Markdown):
- Executive summary (max 5 bullets)
- Key provisions / concepts (bullets)
- Caveat: one sentence stating this is not legal advice ({caveat})
"""
    else:
        task_hint = (
            "Provide (1) an executive summary and (2) a plain-language paraphrase of what the retrieved sentences say, "
            "and (optionally) brief relevance notes to the user's query. "
            "Do not provide legal advice and do not speculate beyond the retrieved text."
        )
        per_sentence_line = "- Per-sentence paraphrases (bullets; optional)\n" if wants_per_sentence else ""
        output_format = f"""Output format (Markdown):
- Executive summary (max 5 bullets)
- Key provisions / concepts (bullets)
- Plain-language paraphrase (2–6 sentences)
- Relevance notes to the query (bullets, optional)
{per_sentence_line}- Caveat: one sentence stating this is not legal advice ({caveat})
"""

    grounding_rules = [
        "Use ONLY the retrieved sentences as evidence.",
        "Cite evidence using rank citations like [1], [2] in every bullet.",
        f"If something is not supported by the citations, write: '{unsupported_msg}' and do not add it.",
    ]
    if strict_grounding:
        grounding_rules.append("STRICT MODE: omit any statement that cannot be grounded with at least one citation.")
    if heterogeneous:
        grounding_rules.append(
            "The evidence is heterogeneous; be conservative in connecting it to the query. If linkage is unclear, say it is not supported."
        )
    if sparse:
        grounding_rules.append(
            "The evidence is sparse; relevance notes are optional and should be omitted if not clearly supported."
        )

    system = (
        "You are a compliance-oriented legal text summarization assistant. "
        "You summarize and paraphrase retrieved statutory sentences neutrally and professionally. "
        "You do not give legal advice, you do not speculate beyond the text, and you always cite evidence by rank."
    )

    rules_md = "\n- ".join(grounding_rules)

    user = f"""{lang_hint}
Task: {task_hint}

Retrieval method: {method_label}
User query / context: {query_text.strip()}
Evidence snapshot: {evidence_snapshot}

Grounding rules:
- {rules_md}

Retrieved sentences (each line starts with a rank citation):
{compact_hits}

{output_format}"""

    prompt = f"<|system|>\n{system}\n<|user|>\n{user}\n<|assistant|>\n"

    out = llm(
        prompt,
        max_tokens=int(max_new_tokens),
        temperature=float(temperature),
        top_p=0.9,
        stop=["</s>", "<|user|>", "<|system|>", "<|assistant|>", "<|end|>"],
    )
    return (out.get("choices", [{}])[0].get("text", "") or "").strip()



def generate_retrieval_summary(
    df_hits: pd.DataFrame,
    *,
    query_text: str,
    method_label: str,
    backend: str,
    task: str,
    output_language: str,
    seq2seq_model_id: str,
    gguf_path: str,
    max_new_tokens: int,
    temperature: float,
    n_ctx: int,
    n_threads: int,
    strict_grounding: bool = True,
    min_evidence_sentences: int = 3,
    heterogeneity_threshold: int = 6,
) -> str:
    """Generate a professional, citation-forward summary / paraphrase for a retrieval result table."""
    backend0 = (backend or "off").strip().lower()
    if backend0 in {"off", "none", "disabled"}:
        return (
            "_LLM post-processing is disabled._\n\n"
            "Set **LLM backend** in the sidebar to **Transformers (seq2seq)** or **llama.cpp (GGUF)**, then run retrieval again.\n"
            "If you only want to recompute the summary after changing settings, use the **Generate LLM summary** button in the relevant accordion."
        )
    if df_hits is None or df_hits.empty:
        return ""

    # Evidence snapshot for transparency and hallucination-avoidance
    ev = _evidence_stats(df_hits, max_items=10)
    n_sent = int(ev.get("n_sent", 0.0) or 0)
    n_law = int(ev.get("n_law_types", 0.0) or 0)
    evidence_note = f"**Evidence snapshot:** {n_sent} retrieved sentences; {n_law} distinct law types."
    if heterogeneity_threshold and n_law >= int(heterogeneity_threshold):
        evidence_note += " Evidence is heterogeneous; relevance mapping may be limited."
    if min_evidence_sentences and n_sent < int(min_evidence_sentences):
        evidence_note += " Evidence is sparse; interpret cautiously."

    compact = _hits_to_compact_text(df_hits, max_items=10, max_chars=9000)
    if not compact:
        return ""

    task0 = (task or "").strip().lower()
    wants_per_sentence = ("per-sentence" in task0) or ("per sentence" in task0) or ("paraphrase" in task0)
    per_sentence_only = wants_per_sentence and (("only" in task0) or ("nur" in task0))

    excerpts = _key_excerpts_md(df_hits, max_items=6)

    try:
        # -------- Transformers seq2seq (summary-only) --------
        if backend0.startswith("transformers") or backend0.startswith("seq2seq") or backend0.startswith("mt5"):
            model_id = (seq2seq_model_id or _DEFAULT_SEQ2SEQ_MODEL).strip()
            caveat = (
                "Hinweis: Automatisch generierte Darstellung; keine Rechtsberatung."
                if not (output_language or "").lower().startswith("en")
                else "Note: Automatically generated rendering; not legal advice."
            )
            parts = [evidence_note]
            if per_sentence_only:
                if excerpts:
                    parts.append(f"### Key excerpts ({method_label})\n\n{excerpts}")
                parts.append(f"_{caveat}_")
                return "\n\n".join(parts).strip()

            summ = _seq2seq_summarize(compact, model_id=model_id, max_new_tokens=int(max_new_tokens))
            parts.append(f"### Summary ({method_label})\n\n{summ}")
            if excerpts:
                parts.append(f"### Key excerpts\n\n{excerpts}")
            parts.append(f"_{caveat}_")
            return "\n\n".join(parts).strip()

        # -------- llama.cpp (grounded instruction model) --------
        if backend0.startswith("llama"):
            caveat = (
                "Hinweis: Automatisch generierte Darstellung; keine Rechtsberatung."
                if not (output_language or "").lower().startswith("en")
                else "Note: Automatically generated rendering; not legal advice."
            )
            if not gguf_path or not str(gguf_path).strip():
                parts = [evidence_note, "**LLM backend is set to llama.cpp, but no GGUF model path is configured.**"]
                if excerpts:
                    parts.append(f"### Key excerpts\n\n{excerpts}")
                parts.append(f"_{caveat}_")
                return "\n\n".join(parts).strip()

            out = _llama_cpp_generate(
                compact,
                query_text=query_text,
                method_label=method_label,
                output_language=output_language,
                gguf_path=gguf_path,
                max_new_tokens=int(max_new_tokens),
                temperature=float(temperature),
                n_ctx=int(n_ctx),
                n_threads=int(n_threads),
                task=task,
                strict_grounding=bool(strict_grounding),
                evidence_stats=ev,
                heterogeneity_threshold=int(heterogeneity_threshold),
                min_evidence_sentences=int(min_evidence_sentences),
            )
            return "\n\n".join([evidence_note, out]).strip()

        return f"**LLM post-processing is enabled, but backend '{backend}' is not recognized.**"
    except Exception as e:
        return f"**LLM post-processing error ({method_label}):** {type(e).__name__}: {e}"



def _coerce_hits_df(obj: Any) -> pd.DataFrame:
    """Best-effort conversion of Gradio Dataframe payloads into a pandas DataFrame."""
    if obj is None:
        return pd.DataFrame()
    if isinstance(obj, pd.DataFrame):
        return obj

    # Gradio may pass dict payloads (e.g., {"headers": [...], "data": [...]})
    if isinstance(obj, dict) and ("data" in obj or "headers" in obj):
        data = obj.get("data") or []
        headers = obj.get("headers")
        if headers:
            return pd.DataFrame(data, columns=headers)
        return pd.DataFrame(data)

    # Lists/tuples: list of dicts or list of rows
    if isinstance(obj, (list, tuple)):
        if len(obj) == 0:
            return pd.DataFrame()
        if isinstance(obj[0], dict):
            return pd.DataFrame(list(obj))
        cols = ["rank", "score", "law_type", "page_no", "sentence"]
        if isinstance(obj[0], (list, tuple)) and len(obj[0]) == len(cols):
            return pd.DataFrame(list(obj), columns=cols)
        return pd.DataFrame(list(obj))

    # Fallback: try constructor
    try:
        return pd.DataFrame(obj)
    except Exception:
        return pd.DataFrame()
# ----------------------------- UI callables -----------------------------
def ui_load_report(report_path: str) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, go.Figure, go.Figure, go.Figure, go.Figure]:
    tables = parse_report_tables(report_path)
    if "table1" not in tables:
        raise ValueError("Could not parse Table 1 from report. Ensure it contains the Table 1 markdown.")
    df1 = tables["table1"]
    df2 = tables.get("table2", pd.DataFrame())
    df3 = tables.get("table3", pd.DataFrame())
    df5 = tables.get("table5", pd.DataFrame())

    fig1 = fig_main_metrics(df1)
    fig2 = fig_mean_consensus_fraction(df1)
    fig_lift = fig_mean_lift(df1)
    fig3 = fig_routing(df5) if not df5.empty else go.Figure()

    return df1, df2, df3, fig1, fig2, fig_lift, fig3


@lru_cache(maxsize=2)
def _load_corpora(
    base_dir: str,
    corpus_parquet: str,
    mb_npz: str,
    kahm_npz: str,
) -> Tuple[pd.DataFrame, CorpusIndex, CorpusIndex]:
    df = _load_dataframe(str(Path(base_dir) / corpus_parquet))
    mb_bundle = _load_bundle(str(Path(base_dir) / mb_npz))
    k_bundle = _load_bundle(str(Path(base_dir) / kahm_npz))

    mb_ids, mb_emb = _subset_to_metadata(mb_bundle, df)
    k_ids, k_emb = _subset_to_metadata(k_bundle, df)

    # Normalize (safe even if already normalized)
    if eval3 is None:
        raise ImportError(f"Could not import evaluate_three_embeddings.py: {_EVAL3_IMPORT_ERROR}")
    mb_emb = eval3.l2_normalize_rows(mb_emb)
    k_emb = eval3.l2_normalize_rows(k_emb)

    mb_index = _build_faiss_index(mb_emb)
    k_index = _build_faiss_index(k_emb)

    return (
        df,
        CorpusIndex(name="MB corpus", sentence_ids=mb_ids, embeddings=mb_emb, faiss_index=mb_index),
        CorpusIndex(name="KAHM corpus", sentence_ids=k_ids, embeddings=k_emb, faiss_index=k_index),
    )


@lru_cache(maxsize=1)
def _load_test_queries(module_attr: str) -> Tuple[List[str], List[str]]:
    if eval3 is None:
        raise ImportError(f"Could not import evaluate_three_embeddings.py: {_EVAL3_IMPORT_ERROR}")
    qs = eval3.load_query_set(module_attr)
    texts = eval3.extract_query_texts(qs)
    labels = eval3.extract_consensus_laws(qs)
    return texts, labels


def _build_query_choices(texts: List[str], labels: List[str]) -> List[str]:
    out = []
    for i, (t, lab) in enumerate(zip(texts, labels)):
        t0 = (t or "").replace("\n", " ").strip()
        if len(t0) > 110:
            t0 = t0[:110] + "…"
        lab0 = (lab or "").strip()
        suffix = f" [{lab0}]" if lab0 else ""
        out.append(f"{i:03d}: {t0}{suffix}")
    return out


def ui_retrieve_for_selection(
    selection: str,
    *,
    base_dir: str,
    corpus_parquet: str,
    mb_npz: str,
    kahm_npz: str,
    idf_svd_model: str,
    kahm_model: str,
    kahm_mode: str,
    kahm_batch: int,
    top_k: int,
    query_set: str,
    llm_backend: str,
    llm_task: str,
    llm_output_language: str,
    llm_seq2seq_model_id: str,
    llm_gguf_path: str,
    llm_max_new_tokens: int,
    llm_temperature: float,
    llm_n_ctx: int,
    llm_n_threads: int,
) -> Tuple[str, pd.DataFrame, go.Figure, str, pd.DataFrame, go.Figure, str]:
    texts, labels = _load_test_queries(query_set)
    # Parse index from selection string
    m = re.match(r"^\s*(\d+)\s*:", selection)
    if not m:
        raise ValueError("Could not parse selected query index.")
    idx = int(m.group(1))
    if idx < 0 or idx >= len(texts):
        raise IndexError("Selected query index out of range.")
    query_text = texts[idx]
    label = labels[idx] if idx < len(labels) else ""
    header = f"Selected query #{idx:03d}  |  Gold (human) label: {label}" if label else f"Selected query #{idx:03d}"

    df_meta, mb_corpus, k_corpus = _load_corpora(base_dir, corpus_parquet, mb_npz, kahm_npz)
    q_emb = _embed_query_kahm(
        query_text,
        idf_svd_model_path=str(Path(base_dir) / idf_svd_model),
        kahm_model_path=str(Path(base_dir) / kahm_model),
        kahm_mode=kahm_mode,
        kahm_batch=int(kahm_batch),
    )

    # KAHM(q->MB)
    s1, ids1 = _search_topk(mb_corpus, q_emb, int(top_k))
    out1, fig1 = _format_hits(df_meta, s1, ids1)
    sum1 = generate_retrieval_summary(
        out1,
        query_text=query_text,
        method_label="KAHM(query→MB corpus)",
        backend=llm_backend,
        task=llm_task,
        output_language=llm_output_language,
        seq2seq_model_id=llm_seq2seq_model_id,
        gguf_path=llm_gguf_path,
        max_new_tokens=int(llm_max_new_tokens),
        temperature=float(llm_temperature),
        n_ctx=int(llm_n_ctx),
        n_threads=int(llm_n_threads),
    )

    # Full-KAHM (q->KAHM corpus)
    s2, ids2 = _search_topk(k_corpus, q_emb, int(top_k))
    out2, fig2 = _format_hits(df_meta, s2, ids2)
    sum2 = generate_retrieval_summary(
        out2,
        query_text=query_text,
        method_label="Full-KAHM (query→KAHM corpus)",
        backend=llm_backend,
        task=llm_task,
        output_language=llm_output_language,
        seq2seq_model_id=llm_seq2seq_model_id,
        gguf_path=llm_gguf_path,
        max_new_tokens=int(llm_max_new_tokens),
        temperature=float(llm_temperature),
        n_ctx=int(llm_n_ctx),
        n_threads=int(llm_n_threads),
    )

    return header, out1, fig1, sum1, out2, fig2, sum2





def ui_retrieve_custom(
    query_text: str,
    *,
    base_dir: str,
    corpus_parquet: str,
    mb_npz: str,
    kahm_npz: str,
    idf_svd_model: str,
    kahm_model: str,
    kahm_mode: str,
    kahm_batch: int,
    top_k: int,
    llm_backend: str,
    llm_task: str,
    llm_output_language: str,
    llm_seq2seq_model_id: str,
    llm_gguf_path: str,
    llm_max_new_tokens: int,
    llm_temperature: float,
    llm_n_ctx: int,
    llm_n_threads: int,
) -> Tuple[pd.DataFrame, go.Figure, str, pd.DataFrame, go.Figure, str]:
    if not query_text or not query_text.strip():
        raise ValueError("Please enter a query.")
    df_meta, mb_corpus, k_corpus = _load_corpora(base_dir, corpus_parquet, mb_npz, kahm_npz)
    q_emb = _embed_query_kahm(
        query_text.strip(),
        idf_svd_model_path=str(Path(base_dir) / idf_svd_model),
        kahm_model_path=str(Path(base_dir) / kahm_model),
        kahm_mode=kahm_mode,
        kahm_batch=int(kahm_batch),
    )
    s1, ids1 = _search_topk(mb_corpus, q_emb, int(top_k))
    out1, fig1 = _format_hits(df_meta, s1, ids1)
    sum1 = generate_retrieval_summary(
        out1,
        query_text=query_text,
        method_label="KAHM(query→MB corpus)",
        backend=llm_backend,
        task=llm_task,
        output_language=llm_output_language,
        seq2seq_model_id=llm_seq2seq_model_id,
        gguf_path=llm_gguf_path,
        max_new_tokens=int(llm_max_new_tokens),
        temperature=float(llm_temperature),
        n_ctx=int(llm_n_ctx),
        n_threads=int(llm_n_threads),
    )

    s2, ids2 = _search_topk(k_corpus, q_emb, int(top_k))
    out2, fig2 = _format_hits(df_meta, s2, ids2)
    sum2 = generate_retrieval_summary(
        out2,
        query_text=query_text,
        method_label="Full-KAHM (query→KAHM corpus)",
        backend=llm_backend,
        task=llm_task,
        output_language=llm_output_language,
        seq2seq_model_id=llm_seq2seq_model_id,
        gguf_path=llm_gguf_path,
        max_new_tokens=int(llm_max_new_tokens),
        temperature=float(llm_temperature),
        n_ctx=int(llm_n_ctx),
        n_threads=int(llm_n_threads),
    )

    return out1, fig1, sum1, out2, fig2, sum2







# ----------------------------- Private document (BYO) helpers -----------------------------
# The dashboard can optionally accept a user-uploaded document (txt/md/docx/pdf), extract sentences,
# and use a selected sentence as a query for the same retrieval methods.
try:  # python-docx
    import docx  # type: ignore
except Exception:  # pragma: no cover
    docx = None  # type: ignore


def _read_text_from_path(path: str) -> str:
    """
    Read text from a local file path (txt/md/docx/pdf).
    Returns extracted text as a single string.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")
    ext = p.suffix.lower()

    if ext in {".txt", ".md", ".csv", ".log"}:
        data = p.read_bytes()
        for enc in ("utf-8", "utf-16", "cp1252", "latin-1"):
            try:
                return data.decode(enc)
            except Exception:
                continue
        # last resort
        return data.decode("utf-8", errors="replace")

    if ext == ".docx":
        if docx is None:
            raise ImportError("python-docx is not installed; cannot read .docx files.")
        d = docx.Document(str(p))
        parts = []
        for para in d.paragraphs:
            t = (para.text or "").strip()
            if t:
                parts.append(t)
        # also include table text (common in policies/contracts)
        for tbl in d.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        parts.append(t)
        return "\n".join(parts)

    if ext == ".pdf":
        # Prefer PyMuPDF (fitz); fallback to pypdf if needed.
        try:
            import fitz  # type: ignore
            doc = fitz.open(str(p))
            pages = []
            for i in range(len(doc)):
                pages.append(doc.load_page(i).get_text("text"))
            doc.close()
            return "\n".join(pages)
        except Exception:
            try:
                from pypdf import PdfReader  # type: ignore
                reader = PdfReader(str(p))
                pages = []
                for page in reader.pages:
                    pages.append(page.extract_text() or "")
                return "\n".join(pages)
            except Exception as e:
                raise ImportError(f"Could not read PDF. Install PyMuPDF (fitz) or pypdf. Details: {e}")

    raise ValueError(f"Unsupported document type: {ext}. Please upload a .txt, .md, .docx, or .pdf file.")


_ABBREV_RE = re.compile(
    r"(?:\b(?:z\.B|u\.a|u\.U|bzw|vgl|ggf|usw|etc|e\.g|i\.e|Dr|Prof|Abs|Art|Nr|No|Stk|Zl)\.)$",
    re.IGNORECASE,
)


def _split_sentences(text: str, *, max_sentences: int = 5000) -> List[str]:
    """
    Heuristic sentence splitter (German/English friendly).
    Returns a list of sentences; caps output to max_sentences for UI stability.
    """
    if not text:
        return []
    # normalize whitespace
    t = re.sub(r"\s+", " ", text.replace("\u00ad", " ")).strip()
    if not t:
        return []

    # Candidate split on punctuation followed by whitespace + a likely sentence start.
    parts = re.split(r'(?<=[\.\!\?])["\')\]]*\s+(?=(?:[A-ZÄÖÜ0-9„"\'\(\[]))', t)

    # Post-process: merge obvious false splits (abbreviations, tiny fragments)
    merged: List[str] = []
    for seg in parts:
        seg = seg.strip()
        if not seg:
            continue
        if merged:
            prev = merged[-1]
            if _ABBREV_RE.search(prev) or len(prev) < 25:
                merged[-1] = (prev + " " + seg).strip()
                continue
        merged.append(seg)

    # Fallback: if we failed to split, try line-ish splitting
    if len(merged) <= 1:
        merged = [s.strip() for s in re.split(r"(?:\s*[\n\r]\s*|;\s+)", text) if s.strip()]

    # Cap for UI
    if len(merged) > int(max_sentences):
        merged = merged[: int(max_sentences)]
    return merged


def _build_sentence_choices(sentences: List[str]) -> List[str]:
    out: List[str] = []
    for i, s in enumerate(sentences):
        s0 = (s or "").replace("\n", " ").strip()
        if len(s0) > 120:
            s0 = s0[:120] + "…"
        out.append(f"{i:04d}: {s0}")
    return out


def _empty_retrieval_df() -> pd.DataFrame:
    return pd.DataFrame(columns=["rank", "score", "law_type", "page_no", "sentence"])


def ui_parse_private_document(
    file_path: Optional[str],
) -> Tuple[Dict[str, Any], Any, str, str, str, pd.DataFrame, go.Figure, str, pd.DataFrame, go.Figure, str]:
    """
    Parse an uploaded private document, extract sentences, and prepare UI elements.

    Returns:
      - state dict (filename, sentences)
      - dropdown update (choices + default)
      - info markdown
      - initial selected sentence preview
      - cleared header markdown
      - empty result table + plot (KAHM->MB)
      - empty summary markdown (KAHM->MB)
      - empty result table + plot (Full-KAHM)
      - empty summary markdown (Full-KAHM)
    """
    empty_df = _empty_retrieval_df()
    empty_fig = go.Figure()
    empty_summary = ""

    if not file_path:
        info = "Upload a document to begin. Supported: **.txt, .md, .docx, .pdf**."
        return (
            {"filename": "", "sentences": []},
            gr.update(choices=[], value=None),
            info,
            "",
            "",
            empty_df,
            empty_fig,
            empty_summary,
            empty_df,
            empty_fig,
            empty_summary,
        )

    try:
        text = _read_text_from_path(str(file_path))
        sentences = _split_sentences(text, max_sentences=5000)
        fname = Path(str(file_path)).name
        if not sentences:
            info = f"Loaded **{fname}**, but could not extract any sentences. Try a different file format or a text-based PDF."
            return (
                {"filename": fname, "sentences": []},
                gr.update(choices=[], value=None),
                info,
                "",
                "",
                empty_df,
                empty_fig,
                empty_summary,
                empty_df,
                empty_fig,
                empty_summary,
            )

        choices = _build_sentence_choices(sentences)
        default = choices[0] if choices else None
        preview = sentences[0] if sentences else ""
        info = (
            f"**Loaded:** {fname}\n\n"
            f"- Extracted characters: {len(text):,}\n"
            f"- Extracted sentences: {len(sentences):,} (UI cap: 5,000)\n\n"
            "Select a sentence below and run retrieval."
        )
        state = {"filename": fname, "sentences": sentences}
        return (
            state,
            gr.update(choices=choices, value=default),
            info,
            preview,
            "",
            empty_df,
            empty_fig,
            empty_summary,
            empty_df,
            empty_fig,
            empty_summary,
        )
    except Exception as e:
        info = f"Could not read the uploaded document. Details: {type(e).__name__}: {e}"
        return (
            {"filename": "", "sentences": []},
            gr.update(choices=[], value=None),
            info,
            "",
            "",
            empty_df,
            empty_fig,
            empty_summary,
            empty_df,
            empty_fig,
            empty_summary,
        )





def ui_private_sentence_preview(selection: str, state: Dict[str, Any]) -> str:
    if not selection or not state or not state.get("sentences"):
        return ""
    m = re.match(r"^\s*(\d+)\s*:", str(selection))
    if not m:
        return ""
    idx = int(m.group(1))
    sents = state.get("sentences", [])
    if idx < 0 or idx >= len(sents):
        return ""
    return str(sents[idx])


def ui_retrieve_private_sentence(
    selection: str,
    state: Dict[str, Any],
    *,
    base_dir: str,
    corpus_parquet: str,
    mb_npz: str,
    kahm_npz: str,
    idf_svd_model: str,
    kahm_model: str,
    kahm_mode: str,
    kahm_batch: int,
    top_k: int,
    llm_backend: str,
    llm_task: str,
    llm_output_language: str,
    llm_seq2seq_model_id: str,
    llm_gguf_path: str,
    llm_max_new_tokens: int,
    llm_temperature: float,
    llm_n_ctx: int,
    llm_n_threads: int,
) -> Tuple[str, pd.DataFrame, go.Figure, str, pd.DataFrame, go.Figure, str]:
    if not state or not state.get("sentences"):
        raise ValueError("Please upload a document and select a sentence first.")
    m = re.match(r"^\s*(\d+)\s*:", str(selection))
    if not m:
        raise ValueError("Please select a sentence from the dropdown.")
    idx = int(m.group(1))
    sents: List[str] = state.get("sentences", [])
    if idx < 0 or idx >= len(sents):
        raise IndexError("Selected sentence index out of range.")
    query_text = str(sents[idx]).strip()
    if not query_text:
        raise ValueError("Selected sentence is empty.")

    header = (
        f"**Document:** {state.get('filename', '')}  \n"
        f"**Selected sentence #{idx:04d}:** {query_text}"
    )

    out1, fig1, sum1, out2, fig2, sum2 = ui_retrieve_custom(
        query_text,
        base_dir=base_dir,
        corpus_parquet=corpus_parquet,
        mb_npz=mb_npz,
        kahm_npz=kahm_npz,
        idf_svd_model=idf_svd_model,
        kahm_model=kahm_model,
        kahm_mode=kahm_mode,
        kahm_batch=int(kahm_batch),
        top_k=int(top_k),
        llm_backend=llm_backend,
        llm_task=llm_task,
        llm_output_language=llm_output_language,
        llm_seq2seq_model_id=llm_seq2seq_model_id,
        llm_gguf_path=llm_gguf_path,
        llm_max_new_tokens=int(llm_max_new_tokens),
        llm_temperature=float(llm_temperature),
        llm_n_ctx=int(llm_n_ctx),
        llm_n_threads=int(llm_n_threads),
    )
    return header, out1, fig1, sum1, out2, fig2, sum2







# ----------------------------- KAHM science narrative (dashboard tab) -----------------------------

SCIENCE_MD = r"""
# Embeddings and methods used in this dashboard

This dashboard compares and operationalizes three representations of Austrian-law sentences:

1. **MB embeddings (Mixedbread baseline)** — a high-quality *neural* sentence embedding model used as a strong retrieval baseline and as a “teacher space”.
2. **IDF–SVD embeddings** — a lightweight, *non-neural* dense representation derived from TF‑IDF and latent semantic analysis (LSA).
3. **KAHM embeddings (MB-approx)** — a KAHM regressor that maps IDF–SVD vectors into the MB embedding space via *geometric topic modeling*.

The evaluation report focuses on **sentence-level evidence retrieval** (Top‑K), not single-label classification.

---

## 1) MB embeddings (Mixedbread)

**What they are**  
The “MB” index is built with a SentenceTransformer model:

- **Model:** `mixedbread-ai/deepset-mxbai-embed-de-large-v1` (default in `build_embedding_index_npz.py`)
- **Granularity:** sentence-level (each RIS sentence becomes one vector)
- **Role in this project:** *reference embedding space* for retrieval quality; also the **regression target** for KAHM training.

**Why this model / why MB at all**  
For Austrian laws, the corpus is predominantly German and includes legal phrasing, citations, and domain-specific vocabulary. A strong German-capable semantic embedding model is an appropriate *upper baseline* for retrieval quality and a useful “teacher space” when we want to approximate neural embeddings with cheaper features.

**Where it is produced**  
`build_embedding_index_npz.py` builds `embedding_index.npz` from `ris_sentences.parquet` using SentenceTransformer encoding and stores metadata such as `model_name`, `created_at_utc`, and a dataset fingerprint alongside the embedding matrix.

---

## 2) IDF–SVD embeddings (TF‑IDF → SVD / LSA)

**What they are**  
“IDF–SVD” is a classic, compute-efficient embedding pipeline:

- **TF‑IDF features** (inverse document frequency weighted term features)
  - word n‑grams: **(1, 3)**
  - optional character n‑grams (“char_wb”): **(3, 6)**
  - large feature budgets (defaults): **600k word features**, **300k char features**
- **TruncatedSVD** (randomized) to obtain a dense vector (default **dim=1024**)
- optional **L2 normalization** (disabled by default)

This is implemented in `build_embedding_index_idf_svd_npz.py` as a scikit‑learn `Pipeline`:
`PreprocessTransformer → FeatureUnion(TF‑IDF word + TF‑IDF char) → AdaptiveTruncatedSVD → (optional) Normalizer`.

**Why it was considered**  
IDF–SVD is attractive as a “no-neural-training” baseline and as an input representation for KAHM because it is:

- **cheap to compute** (CPU-friendly; no GPU required),
- **deterministic and auditable** (important in legal/regulated settings),
- **strong on lexical signals** (citations, terminology, statute references),
- a solid bridge between purely lexical retrieval and fully neural embeddings.

**Where it is produced**  
`build_embedding_index_idf_svd_npz.py` writes both:
- `embedding_index_idf_svd.npz` (dense vectors keyed by `sentence_id`), and
- `idf_svd_model.joblib` (portable pipeline for generating query vectors the same way).

---

## 3) KAHM embeddings (geometric topic modeling in embedding space)

**What they are**  
“KAHM embeddings” in this repo are **MB-approx embeddings**: vectors predicted by a KAHM regressor that maps
IDF–SVD inputs into the MB target space.

Training is described in `train_kahm_embedding_regressor.py`:

- **Inputs (X):** L2-normalized IDF–SVD vectors
- **Targets (Y):** L2-normalized MB vectors
- **Objective:** approximate Y from X with a KAHM-style, geometry-driven model.

**Geometrical modeling of topics**  
KAHM operationalizes “topics” as **regions in the target embedding space**:

1. **Cluster the target embeddings (Y)** into `n_clusters` (KMeans / MiniBatchKMeans).  
   Each cluster corresponds to a coherent region of semantic space (a “topic region” in practice).
2. **Learn a geometric descriptor per region** by training one autoencoder per cluster on the *input* vectors that map to that cluster.  
   At inference time, each autoencoder yields a **reconstruction-distance** for the query/input vector, i.e., “how well does this input fit region c?”.
3. **Convert distances into routing weights** (hard or soft routing) and **reconstruct an output embedding** as a weighted combination of cluster centers in the MB space.

This yields an MB-like embedding without running a large neural encoder at query time.

**Two retrieval modes in the dashboard**  

- **KAHM(query→MB corpus):** regress the query into MB space and search the **MB corpus index**.  
  This isolates the benefit of the *query-side* KAHM mapping while keeping the corpus in the original MB space.
- **Full‑KAHM (query→KAHM corpus):** regress the query and search a **KAHM-transformed corpus index**.  
  This evaluates the fully “compressed” pipeline where neither queries nor corpus require MB encoding at runtime.

---

## Practical interpretation

- If **MB > KAHM(query→MB corpus)**, the gap is mostly due to *approximation quality of the query mapping*.
- If **KAHM(query→MB corpus) ≈ Full‑KAHM**, then the corpus transformation is not the limiting factor.
- If **IDF–SVD** is competitive on a subset of queries, those are likely **lexically dominated** information needs (citations, exact phrases).

---

## References

- arXiv: 2512.01025 (advantages motivating KAHM-driven language models)
- See the project report tab for empirical retrieval results on Austrian-law queries.
"""



EMBEDDINGS_QUICKREF_MD = r"""
### Embeddings in this dashboard (quick reference)

- **MB (Mixedbread) embeddings:** SentenceTransformer model `mixedbread-ai/deepset-mxbai-embed-de-large-v1` (strong neural baseline; also the **target** space for KAHM).
- **IDF–SVD embeddings:** TF‑IDF (word 1–3 grams + optional char 3–6 grams) → TruncatedSVD/LSA (**dim=1024** by default). Fast, deterministic, audit-friendly baseline and KAHM input.
- **KAHM embeddings:** KAHM regressor that maps IDF–SVD → MB space via **output-space clustering (“topic regions”)** and **distance-based routing** (hard/soft) to reconstruct MB-like vectors.
"""



# ----------------------------- Management pitch (dashboard tab) -----------------------------
PITCH_MD = r"""
# Management pitch: funding KAHM-driven retrieval and KAHM-efficient seq2seq summarization

## Executive summary
KAHM has demonstrated that we can **replace transformer-based query embedding at retrieval time** with a CPU‑friendly mapping from classic IR features into a strong neural embedding space.

In the latest evaluation on Austrian laws (**200 human‑labeled queries**, **71,069 aligned sentences**, **k=10**), **KAHM(query→MB corpus)** achieved:
- **Hit@10:** 0.895 (0.850, 0.935)
- **MRR@10 (unique laws):** 0.729 (0.679, 0.779)
- **Top‑1 accuracy:** 0.605 (0.535, 0.675)

This represents a clear lift over the low‑cost baseline (**IDF–SVD**) and is statistically *not resolved* from the Mixedbread neural baseline under paired bootstrap on these 200 queries (the 95% CIs for deltas vs Mixedbread include 0; this is not a formal equivalence claim).

**Funding ask:** support a focused program to productize the retrieval stack and extend KAHM from *embedding approximation* into a **computation‑efficient alternative to transformer encoder components** for seq2seq summarization (and adjacent generation/control tasks), using KAHM’s geometry‑based routing as a drop‑in “encoder/selector” that avoids full attention‑heavy encoding.

---

## What KAHM is buying us (in plain business terms)
- **Lower inference cost and latency at scale:** avoid running a large neural encoder for every user query while retaining MB‑level retrieval quality.
- **Deployment practicality:** IDF–SVD feature extraction plus a small KAHM regressor is CPU‑friendly, auditable, and easier to operationalize in constrained environments.
- **Leverage existing indices:** KAHM maps queries into the fixed Mixedbread (MB) embedding space, enabling reuse of an already‑built neural corpus index.

---

## Evidence of traction: Austrian‑law retrieval (what we can credibly claim today)
**1) Strong lift versus the non‑neural baseline**
- Paired deltas (KAHM − IDF–SVD): **Hit@10 +0.190** (+0.130, +0.250), **MRR@10 +0.241** (+0.182, +0.301), **Top‑1 +0.240** (+0.160, +0.320).

**2) Competitive with the neural baseline for retrieval**
- Mixedbread (true) Hit@10: 0.900 (0.855, 0.940) versus KAHM(query→MB corpus) 0.895 (0.850, 0.935).
- Paired deltas versus Mixedbread are small and not resolved under this bootstrap.

**3) Geometry alignment supports generalization**
- Full‑KAHM embeddings show high cosine alignment with Mixedbread geometry (mean corpus cosine ≈ 0.9405), and recover similar **law‑level** neighborhoods.

Operationally, this means the dashboard’s retrieval results are not a demo artifact: KAHM is learning a stable mapping into a high‑quality semantic space that is already useful for legal evidence retrieval.

---

## How KAHM enables a compute‑efficient alternative to transformer seq2seq components
The current dashboard uses a lightweight transformer seq2seq model to summarize retrieved evidence. This is effective, but the **transformer encoder** is still the dominant inference cost when you scale to many queries or longer inputs.

KAHM offers a credible path to reduce or remove that cost by replacing attention‑heavy encoding with **(1) cheap features, (2) geometry‑based routing, and (3) retrieval‑first compression**.

### A) “KAHM as encoder replacement” for summarization
1. **Segment the input** (query + retrieved evidence) into sentences/paragraphs.
2. Compute **IDF–SVD** features per segment (CPU‑efficient, deterministic).
3. Apply **KAHM** to map each segment into an MB‑like semantic vector.
4. Use those vectors to produce a **summary plan** (topic distribution + top evidence segments) via routing weights and nearest‑neighbor structure.
5. Generate the final summary with a **small decoder** (small transformer decoder, or a non‑transformer sequence model) conditioned on the plan and the short selected evidence.

Key point: the expensive “encode everything with self‑attention” step is replaced by linear algebra + small models and aggressive evidence selection.

### B) “KAHM‑routed cross‑attention” (hybrid)
Instead of cross‑attending over all tokens/segments, the decoder attends to:
- a small set of **topic region prototypes** (cluster centers in MB space), and
- a small set of **top retrieved evidence segments** per topic.

This approximates the role of attention with **bounded context** and **explicit routing**, reducing compute and making the system easier to govern.

### C) Immediate product benefit: retrieval‑first summarization
Even before a new seq2seq architecture is complete, KAHM improves the end‑to‑end system by:
- reducing the cost of query embedding at retrieval time,
- enabling high‑quality evidence selection (smaller context windows for any summarizer), and
- providing confidence signals (vote margin/entropy) to trigger “abstain/escalate” behavior.

---

## Proposed funded pilot (12–16 weeks) and deliverables
1) **End‑to‑end costed baseline (Weeks 1–3):**
   - Measure latency/cost for (a) current transformer‑encoder summarization and (b) retrieval‑first summarization with KAHM.
   - Establish a target: quality parity within an agreed tolerance at materially lower CPU/GPU cost.

2) **KAHM‑efficient summarizer prototype (Weeks 3–10):**
   - Build the “KAHM encoder replacement” pipeline: KAHM embeddings → evidence selection + plan → small decoder.
   - Evaluate on internal summarization tasks (legal answers, case summaries, compliance notes) with human review.

3) **Hybrid KAHM‑routed attention prototype (Weeks 8–16):**
   - Research prototype using topic prototypes + bounded evidence sets for cross‑attention.
   - Deliver an engineering feasibility report and clear go/no‑go criteria.

4) **Packaging for product and IP (Weeks 10–16):**
   - A production‑ready retrieval module (KAHM query embedding + FAISS search + confidence routing).
   - Evaluation harness and monitoring metrics.
   - Identify protectable components (routing, planning, bounded‑attention interfaces).

---

## Lean resource request
- **Staffing:** 1 research engineer (FT) + 0.2–0.4 legal/domain SME (PT) + light security/compliance review.
- **Compute:** CPU‑first for KAHM; small GPU budget for baseline comparisons and decoder experiments.

## Decision criteria (what “success” looks like)
- A demonstrable **KAHM‑efficient summarization** prototype that achieves acceptable quality at materially lower incremental compute/latency than transformer‑encoder seq2seq.
- A quantified business case (cost, latency, governance) and a go/no‑go recommendation for productization.
- Defensible technical assets: benchmarks, evaluation suite, and prototype code that can be leveraged as IP.
"""


# ----------------------------- Build UI -----------------------------
CSS = r""".gradio-container {font-family: ui-sans-serif, system-ui, -apple-system, Segoe UI, Roboto, Helvetica, Arial, "Apple Color Emoji", "Segoe UI Emoji";}
#title {margin-bottom: 0.2rem;}
.small-note {color: #666; font-size: 0.9rem;}
/* Layout tuning */
.config-col {max-width: 360px;}
.results-col {min-width: 720px;}
/* Make retrieval tables feel less cramped */
.kahm-table .gr-dataframe {min-height: 260px;}
/* Widen retrieval table key columns */
.retrieval-table table th:nth-child(1),
.retrieval-table table td:nth-child(1) {min-width: 56px; width: 56px;}
.retrieval-table table th:nth-child(2),
.retrieval-table table td:nth-child(2) {min-width: 96px; width: 96px;}
/* Give the sentence column more room */
.retrieval-table table th:nth-child(5),
.retrieval-table table td:nth-child(5) {min-width: 560px;}
"""


def build_app() -> gr.Blocks:
    """Build and return the Gradio dashboard.

    CSS handling is version-tolerant: newer Gradio versions accept `css=` in Blocks;
    older versions will ignore it via the fallback path.
    """
    try:
        demo = gr.Blocks(title="KAHM Embeddings Dashboard", css=CSS)
    except TypeError:  # pragma: no cover
        demo = gr.Blocks(title="KAHM Embeddings Dashboard")

    with demo:
        gr.Markdown("# KAHM embeddings dashboard", elem_id="title")
        gr.Markdown(
            "Dashboard for (1) visualizing the evaluation report and (2) interactively retrieving evidence sentences "
            "with **KAHM(query→MB corpus)** and **Full-KAHM**.",
            elem_classes=["small-note"],
        )

        with gr.Row():
            with gr.Column(scale=1, min_width=320, elem_classes=['config-col']):
                gr.Markdown("## Configuration")
                with gr.Accordion('Configuration settings', open=True):
                    base_dir = gr.Textbox(value=".", label="Project/data directory (base_dir)")
                    report_path = gr.Textbox(value="kahm_evaluation_report.md", label="Report markdown path (relative to base_dir allowed)")
                    corpus_parquet = gr.Textbox(value="ris_sentences.parquet", label="Corpus parquet (sentence metadata)")
                    mb_npz = gr.Textbox(value="embedding_index.npz", label="Mixedbread corpus embeddings (.npz)")
                    kahm_npz = gr.Textbox(value="embedding_index_kahm_mixedbread_approx.npz", label="Full-KAHM corpus embeddings (.npz)")
                    idf_svd_model = gr.Textbox(value="idf_svd_model.joblib", label="IDF–SVD model (joblib)")
                    kahm_model = gr.Textbox(value="kahm_regressor_idf_to_mixedbread.joblib", label="KAHM regressor (joblib)")
                    kahm_mode = gr.Dropdown(choices=["soft", "hard"], value="soft", label="KAHM mode")
                    kahm_batch = gr.Number(value=512, label="KAHM batch size", precision=0)
                    top_k = gr.Slider(minimum=1, maximum=50, step=1, value=10, label="Top-K for retrieval")
                    query_set = gr.Textbox(value="query_set.TEST_QUERY_SET", label="Query set (module.attr)")


                    with gr.Accordion("LLM post-processing (optional)", open=False):
                        gr.Markdown(
                            "Summarize and/or paraphrase retrieved sentences for professional presentation. "
                            "This is optional and can run fully locally.\n\n"
                            "- **Transformers (seq2seq)**: small summarization model (CPU-friendly).\n"
                            "- **llama.cpp (GGUF)**: local instruction model for richer, citation-grounded interpretation."
                        )
                        llm_backend = gr.Dropdown(
                            choices=[
                                "Off",
                                "Transformers (seq2seq)",
                                "llama.cpp (GGUF)",
                            ],
                            value="Off",
                            label="LLM backend",
                        )
                        llm_task = gr.Dropdown(
                            choices=["Summary", "Summary + plain-language interpretation"],
                            value="Summary + plain-language interpretation",
                            label="LLM task",
                        )
                        llm_output_language = gr.Dropdown(
                            choices=["German", "English"],
                            value="German",
                            label="Output language",
                        )
                        llm_seq2seq_model_id = gr.Textbox(
                            value=_DEFAULT_SEQ2SEQ_MODEL,
                            label="Transformers model id (seq2seq)",
                        )
                        llm_gguf_path = gr.Textbox(
                            value="",
                            label="GGUF model path (llama.cpp)",
                            placeholder="/path/to/model.gguf",
                        )
                        llm_max_new_tokens = gr.Slider(
                            minimum=64,
                            maximum=512,
                            step=8,
                            value=224,
                            label="Max new tokens (summary)",
                        )
                        llm_temperature = gr.Slider(
                            minimum=0.0,
                            maximum=1.0,
                            step=0.05,
                            value=0.1,
                            label="Temperature (llama.cpp)",
                        )
                        llm_n_ctx = gr.Slider(
                            minimum=1024,
                            maximum=8192,
                            step=256,
                            value=4096,
                            label="Context window n_ctx (llama.cpp)",
                        )
                        llm_n_threads = gr.Number(
                            value=0,
                            label="Threads (llama.cpp, 0=auto)",
                            precision=0,
                        )
                    btn_load_report = gr.Button("Load/refresh report visuals", variant="primary")

                    with gr.Accordion("Embeddings & methods (overview)", open=False):
                        gr.Markdown(EMBEDDINGS_QUICKREF_MD)

            with gr.Column(scale=4, min_width=780, elem_classes=['results-col']):
                with gr.Tabs():
                    with gr.Tab("Results (from report)"):
                        gr.Markdown("## Report highlights")
                        gr.Markdown(EMBEDDINGS_QUICKREF_MD, elem_classes=["small-note"])

                        table1 = gr.Dataframe(label="Table 1: Main retrieval metrics", interactive=False, wrap=True)
                        with gr.Row():
                            plot_main = gr.Plot(label="Main metrics")
                        with gr.Row():
                            plot_cons = gr.Plot(label="Mean consensus fraction")
                            plot_lift = gr.Plot(label="Mean lift vs prior")
                        plot_route = gr.Plot(label="Routing tradeoffs")
                        with gr.Accordion("Delta tables (from report)", open=False):
                            table2 = gr.Dataframe(label="Table 2: Paired deltas (KAHM − IDF)", interactive=False, wrap=True)
                            table3 = gr.Dataframe(label="Table 3: Paired deltas (KAHM − Mixedbread)", interactive=False, wrap=True)

                    with gr.Tab("Explore retrieval"):
                        gr.Markdown(
                            "## Interactive retrieval\n"
                            "This view retrieves **Top-K sentences** for a selected test query or an arbitrary query.\n\n"
                            "**Interpretation note:** the goal is to surface *relevant evidence sentences*, which may span multiple law topics; "
                            "this is not a single-label classification task."
                        )
                        gr.Markdown(EMBEDDINGS_QUICKREF_MD, elem_classes=["small-note"])


                        with gr.Accordion("Select from the 200 test queries", open=True):
                            # We'll populate choices on load
                            test_query_dropdown = gr.Dropdown(choices=[], value=None, label="Pick a test query")
                            btn_retrieve_sel = gr.Button("Retrieve for selected test query", variant="secondary")
                            selected_header = gr.Markdown()

                            gr.Markdown("### KAHM(query→MB corpus)")
                            out_kahm_mb = gr.Dataframe(interactive=False, wrap=True, elem_classes=['kahm-table','retrieval-table'])
                            plot_kahm_mb = gr.Plot()

                            with gr.Accordion("LLM summary & plain-language interpretation", open=False):
                                btn_sum_kahm_mb = gr.Button("Generate LLM summary", variant="secondary")
                                summary_kahm_mb = gr.Markdown()


                            gr.Markdown("### Full-KAHM (query→KAHM corpus)")
                            out_full = gr.Dataframe(interactive=False, wrap=True, elem_classes=['kahm-table','retrieval-table'])
                            plot_full = gr.Plot()

                            with gr.Accordion("LLM summary & plain-language interpretation", open=False):
                                btn_sum_full = gr.Button("Generate LLM summary", variant="secondary")
                                summary_full = gr.Markdown()


                        with gr.Accordion("Search with a custom query", open=False):
                            custom_query = gr.Textbox(lines=2, label="Custom query (German or English)")
                            btn_retrieve_custom = gr.Button("Search", variant="primary")
                            gr.Markdown("### KAHM(query→MB corpus)")
                            out_kahm_mb2 = gr.Dataframe(interactive=False, wrap=True, elem_classes=['kahm-table','retrieval-table'])
                            plot_kahm_mb2 = gr.Plot()

                            with gr.Accordion("LLM summary & plain-language interpretation", open=False):
                                btn_sum_kahm_mb2 = gr.Button("Generate LLM summary", variant="secondary")
                                summary_kahm_mb2 = gr.Markdown()


                            gr.Markdown("### Full-KAHM (query→KAHM corpus)")
                            out_full2 = gr.Dataframe(interactive=False, wrap=True, elem_classes=['kahm-table','retrieval-table'])
                            plot_full2 = gr.Plot()

                            with gr.Accordion("LLM summary & plain-language interpretation", open=False):
                                btn_sum_full2 = gr.Button("Generate LLM summary", variant="secondary")
                                summary_full2 = gr.Markdown()



                    with gr.Tab("Private document"):
                        gr.Markdown(
                            "## Private document retrieval\n"
                            "Upload an internal document, select a sentence, and retrieve related Austrian‑law evidence sentences "
                            "using the same methods as the other retrieval views."
                        )
                        gr.Markdown(
                            "**Privacy note:** the uploaded file is processed by this dashboard instance (your local runtime) and is not "
                            "sent to external services by this code.",
                            elem_classes=["small-note"],
                        )

                        private_doc_state = gr.State({"filename": "", "sentences": []})

                        with gr.Row():
                            with gr.Column(scale=1, min_width=360):
                                private_file = gr.File(
                                    label="1) Upload your document",
                                    file_types=[".txt", ".md", ".docx", ".pdf"],
                                    type="filepath",
                                )
                                private_info = gr.Markdown(elem_classes=["small-note"])
                                private_sentence_dropdown = gr.Dropdown(
                                    choices=[],
                                    value=None,
                                    label="2) Select a sentence from your document",
                                )
                                private_sentence_preview = gr.Textbox(
                                    label="Selected sentence (used as query)",
                                    lines=3,
                                    interactive=False,
                                )
                                private_btn_retrieve = gr.Button(
                                    "3) Retrieve related Austrian‑law sentences",
                                    variant="primary",
                                )

                            with gr.Column(scale=2, min_width=560):
                                private_header = gr.Markdown()

                                gr.Markdown("### KAHM(query→MB corpus)")
                                private_out_kahm = gr.Dataframe(
                                    interactive=False,
                                    wrap=True,
                                    elem_classes=["kahm-table", "retrieval-table"],
                                )
                                private_plot_kahm = gr.Plot()

                                with gr.Accordion("LLM summary & plain-language interpretation", open=False):
                                    btn_sum_private_kahm = gr.Button("Generate LLM summary", variant="secondary")
                                    private_summary_kahm = gr.Markdown()

                                gr.Markdown("### Full-KAHM (query→KAHM corpus)")
                                private_out_full_doc = gr.Dataframe(
                                    interactive=False,
                                    wrap=True,
                                    elem_classes=["kahm-table", "retrieval-table"],
                                )
                                private_plot_full_doc = gr.Plot()

                                with gr.Accordion("LLM summary & plain-language interpretation", open=False):
                                    btn_sum_private_full = gr.Button("Generate LLM summary", variant="secondary")
                                    private_summary_full = gr.Markdown()

                    with gr.Tab("Funding pitch"):
                        gr.Markdown(PITCH_MD)

                    with gr.Tab("KAHM science"):
                        gr.Markdown(SCIENCE_MD)

        # ---------- Wiring ----------
        def _resolve_path(base: str, p: str) -> str:
            pth = Path(p)
            if pth.is_absolute():
                return str(pth)
            return str(Path(base) / p)

        def _load_report_wrapper(
            base: str,
            rp: str,
        ):
            rp2 = _resolve_path(base, rp)
            return ui_load_report(rp2)

        btn_load_report.click(
            fn=_load_report_wrapper,
            inputs=[base_dir, report_path],
            outputs=[table1, table2, table3, plot_main, plot_cons, plot_lift, plot_route],
        )

        # Populate test query dropdown on load
        def _init_queries(qset: str):
            texts, labels = _load_test_queries(qset)
            choices = _build_query_choices(texts, labels)
            default = choices[0] if choices else None
            return gr.update(choices=choices, value=default)

        demo.load(fn=_init_queries, inputs=[query_set], outputs=[test_query_dropdown])

        # Refresh test queries if the query-set module path changes
        query_set.change(fn=_init_queries, inputs=[query_set], outputs=[test_query_dropdown])

        # Retrieve for selected test query
        def _retrieve_sel(
            sel: str,
            base: str,
            corpus_pq: str,
            mb: str,
            kahm_npz_p: str,
            idf_model: str,
            kahm_model_p: str,
            mode: str,
            batch: float,
            k: int,
            qset: str,
            llm_b: str,
            llm_t: str,
            llm_lang: str,
            llm_seq: str,
            llm_gguf: str,
            llm_max_tok: int,
            llm_temp: float,
            llm_ctx: int,
            llm_thr: float,
        ):
            return ui_retrieve_for_selection(
                sel,
                base_dir=base,
                corpus_parquet=corpus_pq,
                mb_npz=mb,
                kahm_npz=kahm_npz_p,
                idf_svd_model=idf_model,
                kahm_model=kahm_model_p,
                kahm_mode=mode,
                kahm_batch=int(batch),
                top_k=int(k),
                query_set=qset,
                llm_backend=llm_b,
                llm_task=llm_t,
                llm_output_language=llm_lang,
                llm_seq2seq_model_id=llm_seq,
                llm_gguf_path=llm_gguf,
                llm_max_new_tokens=int(llm_max_tok),
                llm_temperature=float(llm_temp),
                llm_n_ctx=int(llm_ctx),
                llm_n_threads=int(llm_thr),
            )

        btn_retrieve_sel.click(
            fn=_retrieve_sel,
            inputs=[
                test_query_dropdown,
                base_dir,
                corpus_parquet,
                mb_npz,
                kahm_npz,
                idf_svd_model,
                kahm_model,
                kahm_mode,
                kahm_batch,
                top_k,
                query_set,
                llm_backend,
                llm_task,
                llm_output_language,
                llm_seq2seq_model_id,
                llm_gguf_path,
                llm_max_new_tokens,
                llm_temperature,
                llm_n_ctx,
                llm_n_threads,
            ],
            outputs=[selected_header, out_kahm_mb, plot_kahm_mb, summary_kahm_mb, out_full, plot_full, summary_full],
        )


        # --- On-demand LLM summaries (selected test query) ---
        def _summarize_sel_from_hits(
            sel: str,
            qset: str,
            hits: Any,
            method_label: str,
            llm_b: str,
            llm_t: str,
            llm_lang: str,
            llm_seq: str,
            llm_gguf: str,
            llm_max_tok: int,
            llm_temp: float,
            llm_ctx: int,
            llm_thr: float,
        ) -> str:
            # Resolve the full query text from the selection.
            texts, _labels = _load_test_queries(qset)
            m = re.match(r"^\s*(\d+)\s*:", sel or "")
            if not m:
                return ""
            idx0 = int(m.group(1))
            if idx0 < 0 or idx0 >= len(texts):
                return ""
            query_text = texts[idx0]
            df_hits = _coerce_hits_df(hits)
            return generate_retrieval_summary(
                df_hits,
                query_text=query_text,
                method_label=method_label,
                backend=llm_b,
                task=llm_t,
                output_language=llm_lang,
                seq2seq_model_id=llm_seq,
                gguf_path=llm_gguf,
                max_new_tokens=int(llm_max_tok),
                temperature=float(llm_temp),
                n_ctx=int(llm_ctx),
                n_threads=int(llm_thr),
            )

        btn_sum_kahm_mb.click(
            fn=lambda sel, qset, hits, llm_b, llm_t, llm_lang, llm_seq, llm_gguf, llm_max_tok, llm_temp, llm_ctx, llm_thr: _summarize_sel_from_hits(
                sel,
                qset,
                hits,
                "KAHM(query→MB corpus)",
                llm_b,
                llm_t,
                llm_lang,
                llm_seq,
                llm_gguf,
                llm_max_tok,
                llm_temp,
                llm_ctx,
                llm_thr,
            ),
            inputs=[
                test_query_dropdown,
                query_set,
                out_kahm_mb,
                llm_backend,
                llm_task,
                llm_output_language,
                llm_seq2seq_model_id,
                llm_gguf_path,
                llm_max_new_tokens,
                llm_temperature,
                llm_n_ctx,
                llm_n_threads,
            ],
            outputs=[summary_kahm_mb],
        )

        btn_sum_full.click(
            fn=lambda sel, qset, hits, llm_b, llm_t, llm_lang, llm_seq, llm_gguf, llm_max_tok, llm_temp, llm_ctx, llm_thr: _summarize_sel_from_hits(
                sel,
                qset,
                hits,
                "Full-KAHM (query→KAHM corpus)",
                llm_b,
                llm_t,
                llm_lang,
                llm_seq,
                llm_gguf,
                llm_max_tok,
                llm_temp,
                llm_ctx,
                llm_thr,
            ),
            inputs=[
                test_query_dropdown,
                query_set,
                out_full,
                llm_backend,
                llm_task,
                llm_output_language,
                llm_seq2seq_model_id,
                llm_gguf_path,
                llm_max_new_tokens,
                llm_temperature,
                llm_n_ctx,
                llm_n_threads,
            ],
            outputs=[summary_full],
        )

        # Retrieve for custom query
        def _retrieve_custom(
            q: str,
            base: str,
            corpus_pq: str,
            mb: str,
            kahm_npz_p: str,
            idf_model: str,
            kahm_model_p: str,
            mode: str,
            batch: float,
            k: int,
            llm_b: str,
            llm_t: str,
            llm_lang: str,
            llm_seq: str,
            llm_gguf: str,
            llm_max_tok: int,
            llm_temp: float,
            llm_ctx: int,
            llm_thr: float,
        ):
            return ui_retrieve_custom(
                q,
                base_dir=base,
                corpus_parquet=corpus_pq,
                mb_npz=mb,
                kahm_npz=kahm_npz_p,
                idf_svd_model=idf_model,
                kahm_model=kahm_model_p,
                kahm_mode=mode,
                kahm_batch=int(batch),
                top_k=int(k),
                llm_backend=llm_b,
                llm_task=llm_t,
                llm_output_language=llm_lang,
                llm_seq2seq_model_id=llm_seq,
                llm_gguf_path=llm_gguf,
                llm_max_new_tokens=int(llm_max_tok),
                llm_temperature=float(llm_temp),
                llm_n_ctx=int(llm_ctx),
                llm_n_threads=int(llm_thr),
            )

        btn_retrieve_custom.click(
            fn=_retrieve_custom,
            inputs=[
                custom_query,
                base_dir,
                corpus_parquet,
                mb_npz,
                kahm_npz,
                idf_svd_model,
                kahm_model,
                kahm_mode,
                kahm_batch,
                top_k,
                llm_backend,
                llm_task,
                llm_output_language,
                llm_seq2seq_model_id,
                llm_gguf_path,
                llm_max_new_tokens,
                llm_temperature,
                llm_n_ctx,
                llm_n_threads,
            ],
            outputs=[out_kahm_mb2, plot_kahm_mb2, summary_kahm_mb2, out_full2, plot_full2, summary_full2],
        )


        # --- On-demand LLM summaries (custom query) ---
        def _summarize_custom_from_hits(
            q: str,
            hits: Any,
            method_label: str,
            llm_b: str,
            llm_t: str,
            llm_lang: str,
            llm_seq: str,
            llm_gguf: str,
            llm_max_tok: int,
            llm_temp: float,
            llm_ctx: int,
            llm_thr: float,
        ) -> str:
            df_hits = _coerce_hits_df(hits)
            return generate_retrieval_summary(
                df_hits,
                query_text=q or "",
                method_label=method_label,
                backend=llm_b,
                task=llm_t,
                output_language=llm_lang,
                seq2seq_model_id=llm_seq,
                gguf_path=llm_gguf,
                max_new_tokens=int(llm_max_tok),
                temperature=float(llm_temp),
                n_ctx=int(llm_ctx),
                n_threads=int(llm_thr),
            )

        btn_sum_kahm_mb2.click(
            fn=lambda q, hits, llm_b, llm_t, llm_lang, llm_seq, llm_gguf, llm_max_tok, llm_temp, llm_ctx, llm_thr: _summarize_custom_from_hits(
                q,
                hits,
                "KAHM(query→MB corpus)",
                llm_b,
                llm_t,
                llm_lang,
                llm_seq,
                llm_gguf,
                llm_max_tok,
                llm_temp,
                llm_ctx,
                llm_thr,
            ),
            inputs=[
                custom_query,
                out_kahm_mb2,
                llm_backend,
                llm_task,
                llm_output_language,
                llm_seq2seq_model_id,
                llm_gguf_path,
                llm_max_new_tokens,
                llm_temperature,
                llm_n_ctx,
                llm_n_threads,
            ],
            outputs=[summary_kahm_mb2],
        )

        btn_sum_full2.click(
            fn=lambda q, hits, llm_b, llm_t, llm_lang, llm_seq, llm_gguf, llm_max_tok, llm_temp, llm_ctx, llm_thr: _summarize_custom_from_hits(
                q,
                hits,
                "Full-KAHM (query→KAHM corpus)",
                llm_b,
                llm_t,
                llm_lang,
                llm_seq,
                llm_gguf,
                llm_max_tok,
                llm_temp,
                llm_ctx,
                llm_thr,
            ),
            inputs=[
                custom_query,
                out_full2,
                llm_backend,
                llm_task,
                llm_output_language,
                llm_seq2seq_model_id,
                llm_gguf_path,
                llm_max_new_tokens,
                llm_temperature,
                llm_n_ctx,
                llm_n_threads,
            ],
            outputs=[summary_full2],
        )

        # ---------- Private document tab ----------
        def _parse_private_doc(fp: str):
            return ui_parse_private_document(fp)

        private_file.change(
            fn=_parse_private_doc,
            inputs=[private_file],
            outputs=[
                private_doc_state,
                private_sentence_dropdown,
                private_info,
                private_sentence_preview,
                private_header,
                private_out_kahm,
                private_plot_kahm,
                private_summary_kahm,
                private_out_full_doc,
                private_plot_full_doc,
                private_summary_full,
            ],
        )

        private_sentence_dropdown.change(
            fn=ui_private_sentence_preview,
            inputs=[private_sentence_dropdown, private_doc_state],
            outputs=[private_sentence_preview],
        )

        def _retrieve_private(
            sel: str,
            st: Dict[str, Any],
            base: str,
            corpus_pq: str,
            mb: str,
            kahm_npz_p: str,
            idf_model: str,
            kahm_model_p: str,
            mode: str,
            batch: float,
            k: int,
            llm_b: str,
            llm_t: str,
            llm_lang: str,
            llm_seq: str,
            llm_gguf: str,
            llm_max_tok: int,
            llm_temp: float,
            llm_ctx: int,
            llm_thr: float,
        ):
            return ui_retrieve_private_sentence(
                sel,
                st,
                base_dir=base,
                corpus_parquet=corpus_pq,
                mb_npz=mb,
                kahm_npz=kahm_npz_p,
                idf_svd_model=idf_model,
                kahm_model=kahm_model_p,
                kahm_mode=mode,
                kahm_batch=int(batch),
                top_k=int(k),
                llm_backend=llm_b,
                llm_task=llm_t,
                llm_output_language=llm_lang,
                llm_seq2seq_model_id=llm_seq,
                llm_gguf_path=llm_gguf,
                llm_max_new_tokens=int(llm_max_tok),
                llm_temperature=float(llm_temp),
                llm_n_ctx=int(llm_ctx),
                llm_n_threads=int(llm_thr),
            )

        private_btn_retrieve.click(
            fn=_retrieve_private,
            inputs=[
                private_sentence_dropdown,
                private_doc_state,
                base_dir,
                corpus_parquet,
                mb_npz,
                kahm_npz,
                idf_svd_model,
                kahm_model,
                kahm_mode,
                kahm_batch,
                top_k,
                llm_backend,
                llm_task,
                llm_output_language,
                llm_seq2seq_model_id,
                llm_gguf_path,
                llm_max_new_tokens,
                llm_temperature,
                llm_n_ctx,
                llm_n_threads,
            ],
            outputs=[private_header, private_out_kahm, private_plot_kahm, private_summary_kahm, private_out_full_doc, private_plot_full_doc, private_summary_full],
        )

        # --- On-demand LLM summaries (private document) ---
        def _summarize_private_from_hits(
            q: str,
            hits: Any,
            method_label: str,
            llm_b: str,
            llm_t: str,
            llm_lang: str,
            llm_seq: str,
            llm_gguf: str,
            llm_max_tok: int,
            llm_temp: float,
            llm_ctx: int,
            llm_thr: float,
        ) -> str:
            df_hits = _coerce_hits_df(hits)
            return generate_retrieval_summary(
                df_hits,
                query_text=q or "",
                method_label=method_label,
                backend=llm_b,
                task=llm_t,
                output_language=llm_lang,
                seq2seq_model_id=llm_seq,
                gguf_path=llm_gguf,
                max_new_tokens=int(llm_max_tok),
                temperature=float(llm_temp),
                n_ctx=int(llm_ctx),
                n_threads=int(llm_thr),
            )

        btn_sum_private_kahm.click(
            fn=lambda q, hits, llm_b, llm_t, llm_lang, llm_seq, llm_gguf, llm_max_tok, llm_temp, llm_ctx, llm_thr: _summarize_private_from_hits(
                q,
                hits,
                "KAHM(query→MB corpus)",
                llm_b,
                llm_t,
                llm_lang,
                llm_seq,
                llm_gguf,
                llm_max_tok,
                llm_temp,
                llm_ctx,
                llm_thr,
            ),
            inputs=[
                private_sentence_preview,
                private_out_kahm,
                llm_backend,
                llm_task,
                llm_output_language,
                llm_seq2seq_model_id,
                llm_gguf_path,
                llm_max_new_tokens,
                llm_temperature,
                llm_n_ctx,
                llm_n_threads,
            ],
            outputs=[private_summary_kahm],
        )

        btn_sum_private_full.click(
            fn=lambda q, hits, llm_b, llm_t, llm_lang, llm_seq, llm_gguf, llm_max_tok, llm_temp, llm_ctx, llm_thr: _summarize_private_from_hits(
                q,
                hits,
                "Full-KAHM (query→KAHM corpus)",
                llm_b,
                llm_t,
                llm_lang,
                llm_seq,
                llm_gguf,
                llm_max_tok,
                llm_temp,
                llm_ctx,
                llm_thr,
            ),
            inputs=[
                private_sentence_preview,
                private_out_full_doc,
                llm_backend,
                llm_task,
                llm_output_language,
                llm_seq2seq_model_id,
                llm_gguf_path,
                llm_max_new_tokens,
                llm_temperature,
                llm_n_ctx,
                llm_n_threads,
            ],
            outputs=[private_summary_full],
        )




    return demo


if __name__ == "__main__":
    build_app().launch(share=True)
